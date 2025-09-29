import os
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK


def run(cmd: str) -> str:
    completed = subprocess.run(cmd, shell=True, check=True, capture_output=True, text=True)
    return completed.stdout


def add_code_block(doc: Document, title: str, code_text: str):
    if title:
        doc.add_paragraph(title)
    p = doc.add_paragraph()
    r = p.add_run()
    r.font.name = 'Consolas'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    r.font.size = Pt(9)
    r.add_text(code_text)


def main():
    repo_root = Path(__file__).resolve().parents[1]
    os.chdir(repo_root)

    # Collect metadata
    head = run('git rev-parse HEAD').strip()
    prev = run('git rev-parse HEAD^').strip()
    remote = run('git remote get-url origin').strip()

    # File paths
    app_py = 'app.py'

    # Extract vulnerable snippets from previous commit
    vuln_check_login = run(f"git show {prev}:{app_py} | nl -ba | sed -n '35,55p'")
    vuln_admin_route = run(f"git show {prev}:{app_py} | nl -ba | sed -n '472,492p'")

    # Extract fixed snippets from HEAD
    fix_check_login = run(f"nl -ba {app_py} | sed -n '40,47p'")
    fix_admin_route = run(f"nl -ba {app_py} | sed -n '478,486p'")

    # Build links (best-effort, GitHub blob URLs)
    def blob_url(commit: str, path: str, start: int, end: int) -> str:
        base = remote
        if base.endswith('.git'):
            base = base[:-4]
        return f"{base}/blob/{commit}/{path}#L{start}-L{end}"

    link_vuln_check_login = blob_url(prev, app_py, 40, 45)
    link_vuln_admin_route = blob_url(prev, app_py, 478, 486)
    link_fix_check_login = blob_url(head, app_py, 40, 45)
    link_fix_admin_route = blob_url(head, app_py, 478, 486)

    # Create document
    doc = Document()

    doc.add_heading('OWASP Top 10 Toepassing – Werkplaats (WP2)', level=1)
    doc.add_paragraph('Student: (vul naam/studentnummer in)')
    doc.add_paragraph('Datum: (vul datum in)')

    doc.add_heading('1) Kloon de repository', level=2)
    doc.add_paragraph('Gekloonde repository:')
    doc.add_paragraph(remote)

    doc.add_heading('2) Reproduceerbare kwetsbaarheid', level=2)
    doc.add_paragraph('Kwetsbaarheid gereproduceerd door in te loggen als niet-admin en een admin-gebruiker aan te maken via de route voor nieuwe redacteuren.')
    doc.add_paragraph('Bewijs: SQL toont dat gebruiker attacker@example.com met is_admin=1 is aangemaakt.')

    doc.add_heading('3) OWASP Top 10 categorie', level=2)
    doc.add_paragraph('A01:2021 – Broken Access Control')

    doc.add_heading('4) Kwetsbare code (regels)', level=2)
    doc.add_paragraph(f"Link check_login (kwetsbaar): {link_vuln_check_login}")
    add_code_block(doc, 'check_login (voor fix):', vuln_check_login)
    doc.add_paragraph(f"Link admin-route (kwetsbaar): {link_vuln_admin_route}")
    add_code_block(doc, 'nieuwe_redacteuren (voor fix):', vuln_admin_route)

    doc.add_heading('5) Link naar kwetsbare code (commit)', level=2)
    doc.add_paragraph(f"Commit (kwetsbaar): {prev}")

    doc.add_heading('6) Uitleg waarom kwetsbaar', level=2)
    doc.add_paragraph('De functie check_login controleerde admin-rechten omgekeerd: bij require_admin=True werd juist geblokkeerd wanneer de gebruiker admin was, en toegestaan wanneer de gebruiker geen admin was. Bovendien riep de route voor het aanmaken van een nieuwe redacteur check_login() zonder require_admin=True aan. Hierdoor kon een niet-admin een account met admin-rechten aanmaken (privilege escalation).')

    doc.add_heading('7) Nieuwe code (fix)', level=2)
    doc.add_paragraph(f"Link check_login (fix): {link_fix_check_login}")
    add_code_block(doc, 'check_login (na fix):', fix_check_login)
    doc.add_paragraph(f"Link admin-route (fix): {link_fix_admin_route}")
    add_code_block(doc, 'nieuwe_redacteuren (na fix):', fix_admin_route)

    doc.add_heading('8) Commit ID(s) nieuwe code', level=2)
    doc.add_paragraph(f"Commit (fix): {head}")

    doc.add_heading('9) Waarom dit de kwetsbaarheid verhelpt', level=2)
    doc.add_paragraph('De fix past deny-by-default toe voor admin-acties: wanneer require_admin=True en de sessie niet admin is, volgt een 401. Daarnaast is de admin-route expliciet beschermd met check_login(True). Dit volgt OWASP A01-aanbevelingen: expliciete autorisatie checks op elke gevoelige route, principle of least privilege, en systeematische controle vóór uitvoering van gevoelige acties.')

    doc.add_heading('Gebruik van tools', level=2)
    doc.add_paragraph('Reproductie is geautomatiseerd met curl en verificatie met SQLite queries. (Optioneel) OWASP ZAP of SQLMap inzet kan worden toegevoegd.')

    # Save into docs/
    out_dir = repo_root / 'docs'
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / 'OWASP_WP2_Report.docx'
    doc.save(str(out_path))
    print(f"Report generated: {out_path}")


if __name__ == '__main__':
    main()

